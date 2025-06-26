
#------------------------------------- STREAMLIT PDF TO WORD CONVERTER------------------------------------------#

import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import io

st.title("PDF to Word Converter")

st.write("Upload your pdf here.")
uploaded_file = st.file_uploader(" ", type=['pdf'])

if uploaded_file is not None:
    pdf = PdfReader(uploaded_file)
    st.success("Pdf uploaded successfully")
    #st.write(f"No of pages: {len(pdf.pages)}")

    doc = Document()

    for page_num in range(len(pdf.pages)):
        page = pdf.pages[page_num]
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("[No extractable text found on this page]")

    word_io = io.BytesIO()

    doc.save(word_io)
    word_io.seek(0)

    st.download_button(
        label="Download converted document",
        data=word_io,
        file_name="Coverted.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )












